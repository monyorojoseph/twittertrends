from pickle import NONE
from celery import shared_task
from django.core.mail import EmailMessage
import time, os
from datetime import datetime
from docx import Document
from .config import verify_api


KENYA_WOEID = 23424863


@shared_task(name="ras")
def send_trends():
    
    api = verify_api()
    trend = api.get_place_trends(KENYA_WOEID)
    document = Document()        
    document.add_heading("Kenyan Trends on Twitter", 0)
    document.add_paragraph(f"Most trending hashtag and tweets on Kenyan Twitter. On {datetime.now().strftime('%a %d/%m/%Y')} at {datetime.now().strftime('%H:%M %p')}")
    for tr in trend[0]['trends']:
        document.add_heading(tr['name'], level=3)
        for sr in api.search_tweets(q=f"{tr['query']} -filter:retweets", result_type="mixed", count=10, tweet_mode='extended'):
            p = document.add_paragraph(sr.full_text, style='List Bullet')
            p.add_run(f"@{sr.user.screen_name}").italic = True 
    
    document.save("twitter/kenyantrends.docx")
    
    to = ['monyorojoseph@gmail.com',]
    frm = "shedgehog101@gmail.com"
    subject = f"Twitter Trends In Kenya at {datetime.now().strftime('%H:%M %p')}"
    body = f"Trends at twitter on {datetime.now().strftime('%a %d/%m/%Y')} at {datetime.now().strftime('%H:%M %p')}"
    email = EmailMessage(subject, body, frm, to)
    email.attach_file("twitter/kenyantrends.docx")
    email.send()
    
    time.sleep(100)
    os.remove("twitter/kenyantrends.docx")
    return None

